"""
Reports API endpoints
PDF, Word, Excel export generation
"""

from typing import List, Optional
from datetime import date
from fastapi import APIRouter, Depends, HTTPException, Query, BackgroundTasks
from fastapi.responses import StreamingResponse, FileResponse
from sqlalchemy.ext.asyncio import AsyncSession
import io

from app.core.database import get_async_session
from app.schemas.employee import EmployeeFilter
from app.schemas.analytics import DashboardRequest

router = APIRouter()


# ============= Report Generation Service =============
class ReportService:
    """
    Report generation service for PDF, Word, and Excel exports
    Uses ReportLab, python-docx, and xlsxwriter
    """
    
    def __init__(self, session: AsyncSession):
        self.session = session
    
    async def generate_pdf_report(
        self,
        report_type: str,
        filters: EmployeeFilter,
        sections: List[str],
        include_charts: bool = True
    ) -> bytes:
        """Generate PDF report using ReportLab"""
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import letter, A4
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4)
        styles = getSampleStyleSheet()
        story = []
        
        # Title
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            spaceAfter=30
        )
        story.append(Paragraph(f"HR Analytics Report - {report_type.title()}", title_style))
        story.append(Spacer(1, 12))
        
        # Add sections based on request
        if "executive_summary" in sections:
            story.append(Paragraph("Executive Summary", styles['Heading2']))
            story.append(Paragraph(
                "This report provides comprehensive HR analytics insights...",
                styles['Normal']
            ))
            story.append(Spacer(1, 12))
        
        if "headcount" in sections:
            story.append(Paragraph("Headcount Analysis", styles['Heading2']))
            # Add headcount data table
            data = [
                ['Department', 'Count', 'Percentage'],
                ['Operations', '25,000', '25%'],
                ['Technology', '15,000', '15%'],
                ['Retail Banking', '20,000', '20%'],
            ]
            table = Table(data)
            table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 12),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))
            story.append(table)
            story.append(Spacer(1, 12))
        
        if "compensation" in sections:
            story.append(Paragraph("Compensation Analysis", styles['Heading2']))
            story.append(Paragraph(
                "Salary distribution and pay equity metrics...",
                styles['Normal']
            ))
            story.append(Spacer(1, 12))
        
        if "diversity" in sections:
            story.append(Paragraph("Diversity & Inclusion", styles['Heading2']))
            story.append(Paragraph(
                "Gender distribution and diversity index scores...",
                styles['Normal']
            ))
            story.append(Spacer(1, 12))
        
        if "performance" in sections:
            story.append(Paragraph("Performance Analysis", styles['Heading2']))
            story.append(Paragraph(
                "Performance score distribution and trends...",
                styles['Normal']
            ))
            story.append(Spacer(1, 12))
        
        if "predictions" in sections:
            story.append(Paragraph("AI Predictions & Insights", styles['Heading2']))
            story.append(Paragraph(
                "Attrition risk analysis and performance predictions...",
                styles['Normal']
            ))
            story.append(Spacer(1, 12))
        
        doc.build(story)
        buffer.seek(0)
        return buffer.getvalue()
    
    async def generate_word_report(
        self,
        report_type: str,
        filters: EmployeeFilter,
        sections: List[str]
    ) -> bytes:
        """Generate Word document using python-docx"""
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        doc = Document()
        
        # Title
        title = doc.add_heading(f'HR Analytics Report - {report_type.title()}', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Date
        date_para = doc.add_paragraph(f'Generated: {date.today().isoformat()}')
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        
        # Add sections
        if "executive_summary" in sections:
            doc.add_heading('Executive Summary', level=1)
            doc.add_paragraph(
                'This report provides a comprehensive analysis of HR metrics '
                'and workforce analytics for the organization.'
            )
        
        if "headcount" in sections:
            doc.add_heading('Headcount Analysis', level=1)
            
            # Add table
            table = doc.add_table(rows=4, cols=3)
            table.style = 'Table Grid'
            
            # Header row
            header_cells = table.rows[0].cells
            header_cells[0].text = 'Department'
            header_cells[1].text = 'Count'
            header_cells[2].text = 'Percentage'
            
            # Data rows
            data = [
                ('Operations', '25,000', '25%'),
                ('Technology', '15,000', '15%'),
                ('Retail Banking', '20,000', '20%'),
            ]
            
            for i, row_data in enumerate(data):
                row = table.rows[i + 1].cells
                for j, val in enumerate(row_data):
                    row[j].text = val
            
            doc.add_paragraph()
        
        if "compensation" in sections:
            doc.add_heading('Compensation Analysis', level=1)
            doc.add_paragraph(
                'Analysis of salary distribution, pay equity, and total compensation.'
            )
        
        if "diversity" in sections:
            doc.add_heading('Diversity & Inclusion', level=1)
            doc.add_paragraph(
                'Gender distribution, age demographics, and diversity index scores.'
            )
        
        if "performance" in sections:
            doc.add_heading('Performance Analysis', level=1)
            doc.add_paragraph(
                'Performance score distribution and trend analysis.'
            )
        
        if "predictions" in sections:
            doc.add_heading('AI Predictions & Insights', level=1)
            doc.add_paragraph(
                'Machine learning predictions for attrition, performance, and promotion.'
            )
        
        # Save to buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
    
    async def generate_excel_report(
        self,
        report_type: str,
        filters: EmployeeFilter,
        sheets: List[str]
    ) -> bytes:
        """Generate Excel workbook using xlsxwriter"""
        import xlsxwriter
        
        buffer = io.BytesIO()
        workbook = xlsxwriter.Workbook(buffer, {'in_memory': True})
        
        # Formats
        header_format = workbook.add_format({
            'bold': True,
            'bg_color': '#4472C4',
            'font_color': 'white',
            'border': 1
        })
        
        cell_format = workbook.add_format({
            'border': 1
        })
        
        number_format = workbook.add_format({
            'border': 1,
            'num_format': '#,##0'
        })
        
        percent_format = workbook.add_format({
            'border': 1,
            'num_format': '0.0%'
        })
        
        currency_format = workbook.add_format({
            'border': 1,
            'num_format': 'PKR #,##0'
        })
        
        # Summary sheet
        if "summary" in sheets or len(sheets) == 0:
            ws = workbook.add_worksheet('Summary')
            ws.write('A1', 'HR Analytics Summary', header_format)
            ws.write('A3', 'Total Employees', cell_format)
            ws.write('B3', 100000, number_format)
            ws.write('A4', 'Active Employees', cell_format)
            ws.write('B4', 95000, number_format)
            ws.write('A5', 'Average Salary', cell_format)
            ws.write('B5', 150000, currency_format)
            ws.write('A6', 'Avg Performance', cell_format)
            ws.write('B6', 3.5, cell_format)
            ws.set_column('A:A', 20)
            ws.set_column('B:B', 15)
        
        # Headcount sheet
        if "headcount" in sheets:
            ws = workbook.add_worksheet('Headcount')
            headers = ['Department', 'Count', 'Percentage', 'Avg Tenure']
            for col, header in enumerate(headers):
                ws.write(0, col, header, header_format)
            
            data = [
                ['Operations', 25000, 0.25, 5.2],
                ['Technology', 15000, 0.15, 3.8],
                ['Retail Banking', 20000, 0.20, 6.1],
                ['Corporate Banking', 10000, 0.10, 7.3],
                ['Risk Management', 8000, 0.08, 4.5],
                ['Compliance', 5000, 0.05, 4.2],
                ['HR', 3000, 0.03, 5.0],
                ['Finance', 7000, 0.07, 5.5],
                ['Treasury', 4000, 0.04, 6.8],
                ['Other', 3000, 0.03, 4.0],
            ]
            
            for row, row_data in enumerate(data, 1):
                ws.write(row, 0, row_data[0], cell_format)
                ws.write(row, 1, row_data[1], number_format)
                ws.write(row, 2, row_data[2], percent_format)
                ws.write(row, 3, row_data[3], cell_format)
            
            ws.set_column('A:A', 20)
            ws.set_column('B:D', 12)
        
        # Compensation sheet
        if "compensation" in sheets:
            ws = workbook.add_worksheet('Compensation')
            headers = ['Grade', 'Avg Salary', 'Min', 'Max', 'Count']
            for col, header in enumerate(headers):
                ws.write(0, col, header, header_format)
            
            data = [
                ['OG-4', 50000, 40000, 60000, 15000],
                ['OG-3', 80000, 65000, 95000, 20000],
                ['OG-2', 120000, 100000, 140000, 18000],
                ['OG-1', 180000, 150000, 210000, 15000],
                ['AVP', 280000, 240000, 320000, 12000],
                ['VP', 400000, 350000, 450000, 8000],
                ['SVP', 600000, 500000, 700000, 5000],
                ['EVP', 900000, 750000, 1050000, 3000],
            ]
            
            for row, row_data in enumerate(data, 1):
                ws.write(row, 0, row_data[0], cell_format)
                ws.write(row, 1, row_data[1], currency_format)
                ws.write(row, 2, row_data[2], currency_format)
                ws.write(row, 3, row_data[3], currency_format)
                ws.write(row, 4, row_data[4], number_format)
            
            ws.set_column('A:A', 10)
            ws.set_column('B:E', 15)
        
        # Diversity sheet
        if "diversity" in sheets:
            ws = workbook.add_worksheet('Diversity')
            headers = ['Category', 'Male', 'Female', 'Total', '% Female']
            for col, header in enumerate(headers):
                ws.write(0, col, header, header_format)
            
            data = [
                ['Overall', 70000, 30000, 100000, 0.30],
                ['Leadership', 1800, 200, 2000, 0.10],
                ['Management', 8000, 2000, 10000, 0.20],
                ['Staff', 60200, 27800, 88000, 0.316],
            ]
            
            for row, row_data in enumerate(data, 1):
                ws.write(row, 0, row_data[0], cell_format)
                ws.write(row, 1, row_data[1], number_format)
                ws.write(row, 2, row_data[2], number_format)
                ws.write(row, 3, row_data[3], number_format)
                ws.write(row, 4, row_data[4], percent_format)
            
            ws.set_column('A:A', 15)
            ws.set_column('B:E', 12)
        
        # Employee data sheet
        if "employees" in sheets:
            ws = workbook.add_worksheet('Employee Data')
            headers = ['Employee ID', 'Name', 'Department', 'Grade', 'Salary', 'Performance']
            for col, header in enumerate(headers):
                ws.write(0, col, header, header_format)
            
            # Sample data - in real implementation, this would come from database
            sample_data = [
                ['EMP001', 'Ahmad Khan', 'Operations', 'OG-2', 120000, 4.2],
                ['EMP002', 'Sara Ahmed', 'Technology', 'AVP', 280000, 4.5],
                ['EMP003', 'Hassan Ali', 'Retail Banking', 'OG-1', 180000, 3.8],
            ]
            
            for row, row_data in enumerate(sample_data, 1):
                ws.write(row, 0, row_data[0], cell_format)
                ws.write(row, 1, row_data[1], cell_format)
                ws.write(row, 2, row_data[2], cell_format)
                ws.write(row, 3, row_data[3], cell_format)
                ws.write(row, 4, row_data[4], currency_format)
                ws.write(row, 5, row_data[5], cell_format)
            
            ws.set_column('A:A', 12)
            ws.set_column('B:B', 20)
            ws.set_column('C:C', 15)
            ws.set_column('D:D', 10)
            ws.set_column('E:F', 12)
        
        workbook.close()
        buffer.seek(0)
        return buffer.getvalue()


# ============= API Endpoints =============

@router.post("/pdf")
async def generate_pdf_report(
    report_type: str = Query("comprehensive", description="Report type"),
    departments: Optional[List[str]] = Query(None),
    grades: Optional[List[str]] = Query(None),
    sections: List[str] = Query(
        ["executive_summary", "headcount", "compensation", "diversity", "performance"],
        description="Sections to include"
    ),
    include_charts: bool = Query(True, description="Include charts in report"),
    session: AsyncSession = Depends(get_async_session)
):
    """
    Generate PDF report
    
    Available sections:
    - executive_summary
    - headcount
    - compensation
    - diversity
    - performance
    - attrition
    - predictions
    - recommendations
    """
    filters = EmployeeFilter(
        departments=departments,
        grades=grades
    )
    
    service = ReportService(session)
    pdf_content = await service.generate_pdf_report(
        report_type=report_type,
        filters=filters,
        sections=sections,
        include_charts=include_charts
    )
    
    return StreamingResponse(
        io.BytesIO(pdf_content),
        media_type="application/pdf",
        headers={
            "Content-Disposition": f"attachment; filename=hr_analytics_report.pdf"
        }
    )


@router.post("/word")
async def generate_word_report(
    report_type: str = Query("comprehensive"),
    departments: Optional[List[str]] = Query(None),
    grades: Optional[List[str]] = Query(None),
    sections: List[str] = Query(
        ["executive_summary", "headcount", "compensation", "diversity", "performance"]
    ),
    session: AsyncSession = Depends(get_async_session)
):
    """
    Generate Word document report
    
    Returns .docx file with formatted HR analytics report.
    """
    filters = EmployeeFilter(
        departments=departments,
        grades=grades
    )
    
    service = ReportService(session)
    docx_content = await service.generate_word_report(
        report_type=report_type,
        filters=filters,
        sections=sections
    )
    
    return StreamingResponse(
        io.BytesIO(docx_content),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": f"attachment; filename=hr_analytics_report.docx"
        }
    )


@router.post("/excel")
async def generate_excel_report(
    report_type: str = Query("comprehensive"),
    departments: Optional[List[str]] = Query(None),
    grades: Optional[List[str]] = Query(None),
    sheets: List[str] = Query(
        ["summary", "headcount", "compensation", "diversity", "employees"],
        description="Sheets to include"
    ),
    session: AsyncSession = Depends(get_async_session)
):
    """
    Generate Excel workbook report
    
    Available sheets:
    - summary: Key metrics summary
    - headcount: Headcount by department
    - compensation: Salary analysis
    - diversity: D&I metrics
    - performance: Performance scores
    - employees: Employee data export
    - predictions: AI predictions
    """
    filters = EmployeeFilter(
        departments=departments,
        grades=grades
    )
    
    service = ReportService(session)
    xlsx_content = await service.generate_excel_report(
        report_type=report_type,
        filters=filters,
        sheets=sheets
    )
    
    return StreamingResponse(
        io.BytesIO(xlsx_content),
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={
            "Content-Disposition": f"attachment; filename=hr_analytics_report.xlsx"
        }
    )


@router.post("/scheduled")
async def schedule_report(
    report_type: str,
    format: str = Query("pdf", pattern="^(pdf|word|excel)$"),
    schedule: str = Query("daily", pattern="^(daily|weekly|monthly)$"),
    recipients: List[str] = Query(..., description="Email addresses"),
    departments: Optional[List[str]] = Query(None),
    session: AsyncSession = Depends(get_async_session)
):
    """
    Schedule automated report generation
    
    Reports will be generated and emailed to recipients
    on the specified schedule.
    """
    # In a real implementation, this would:
    # 1. Save schedule to database
    # 2. Set up celery task or cron job
    
    return {
        "message": "Report scheduled successfully",
        "schedule": schedule,
        "format": format,
        "recipients": recipients,
        "report_type": report_type,
        "departments": departments or "All"
    }


@router.get("/templates")
async def list_report_templates():
    """
    List available report templates
    """
    return {
        "templates": [
            {
                "id": "comprehensive",
                "name": "Comprehensive HR Report",
                "description": "Full analytics report with all sections",
                "sections": ["executive_summary", "headcount", "compensation", "diversity", "performance", "attrition", "predictions"]
            },
            {
                "id": "executive",
                "name": "Executive Summary",
                "description": "High-level metrics for leadership",
                "sections": ["executive_summary", "headcount", "compensation"]
            },
            {
                "id": "diversity",
                "name": "Diversity & Inclusion Report",
                "description": "Focus on D&I metrics and analysis",
                "sections": ["executive_summary", "diversity", "recommendations"]
            },
            {
                "id": "compensation",
                "name": "Compensation Analysis",
                "description": "Salary and pay equity analysis",
                "sections": ["executive_summary", "compensation", "salary_bands"]
            },
            {
                "id": "performance",
                "name": "Performance Review",
                "description": "Performance scores and trends",
                "sections": ["executive_summary", "performance", "predictions"]
            },
            {
                "id": "attrition",
                "name": "Attrition Risk Report",
                "description": "Turnover analysis and predictions",
                "sections": ["executive_summary", "attrition", "predictions", "recommendations"]
            }
        ]
    }


@router.post("/custom")
async def generate_custom_report(
    template_id: str,
    format: str = Query("pdf", pattern="^(pdf|word|excel)$"),
    departments: Optional[List[str]] = Query(None),
    grades: Optional[List[str]] = Query(None),
    custom_title: Optional[str] = Query(None),
    session: AsyncSession = Depends(get_async_session)
):
    """
    Generate report from template
    """
    templates = {
        "comprehensive": ["executive_summary", "headcount", "compensation", "diversity", "performance", "attrition", "predictions"],
        "executive": ["executive_summary", "headcount", "compensation"],
        "diversity": ["executive_summary", "diversity"],
        "compensation": ["executive_summary", "compensation"],
        "performance": ["executive_summary", "performance", "predictions"],
        "attrition": ["executive_summary", "attrition", "predictions"]
    }
    
    if template_id not in templates:
        raise HTTPException(
            status_code=400,
            detail=f"Invalid template_id. Available: {list(templates.keys())}"
        )
    
    filters = EmployeeFilter(
        departments=departments,
        grades=grades
    )
    
    service = ReportService(session)
    
    if format == "pdf":
        content = await service.generate_pdf_report(
            report_type=template_id,
            filters=filters,
            sections=templates[template_id]
        )
        media_type = "application/pdf"
        ext = "pdf"
    elif format == "word":
        content = await service.generate_word_report(
            report_type=template_id,
            filters=filters,
            sections=templates[template_id]
        )
        media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ext = "docx"
    else:
        content = await service.generate_excel_report(
            report_type=template_id,
            filters=filters,
            sheets=["summary", "headcount", "compensation", "diversity"]
        )
        media_type = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        ext = "xlsx"
    
    filename = custom_title or f"hr_report_{template_id}"
    
    return StreamingResponse(
        io.BytesIO(content),
        media_type=media_type,
        headers={
            "Content-Disposition": f"attachment; filename={filename}.{ext}"
        }
    )
